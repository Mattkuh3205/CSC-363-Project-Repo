from langchain_core.prompts import ChatPromptTemplate
from langchain_ollama.llms import OllamaLLM
import os
from docx import Document 
from docx2pdf import convert

template = """Question: {question}

Answer: Let's think step by step."""

prompt = ChatPromptTemplate.from_template(template)

model = OllamaLLM(model="llama3.1")

chain = prompt | model

# Generate the Markdown content
answer = chain.invoke({"question": "Create a markdown based resume based on these details: Name: John Doe, Email:john_lol@gmail.com,University: Stanford, Degree: Computer Science, Graduation Year: 2022, Experience: Software Engineer at Google, Skills: Python, Java, C++, Generate a resume that can be used to apply for applications, make the formatting nice and fill in spots that are missing, keep it under 1 page, don't give me a tutorial, actually create this resume for me.FIll it with alot of dummy resume details that you can think of"})

# Write the Markdown content to a file
with open("resume.md", "w") as f:
    f.write(answer)

# Read the Markdown file
with open("resume.md", "r") as file:
    markdown_text = file.read()

# Create a Word document
doc = Document()
doc.add_heading('Resume', level=1)

lines = markdown_text.split('\n')
for line in lines:
    if line.startswith('**') and line.endswith('**'):
        doc.add_heading(line.strip('**'), level=2)
    else:
        doc.add_paragraph(line)

doc.save("resume.docx")

# Convert the Word document to PDF
convert("resume.docx", "resume.pdf")
