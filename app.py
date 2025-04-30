from flask import Flask, render_template, request, send_file
from langchain_openai import ChatOpenAI
from docx import Document 
from docx2pdf import convert
import PyPDF2

import io
import os

app = Flask(__name__)

def read_resume_file(file):
    """Read content from uploaded resume file"""
    if file.filename.endswith('.txt') or file.filename.endswith('.md'):
        return file.read().decode('utf-8')
    elif file.filename.endswith('.docx'):
        doc = Document(file)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    elif file.filename.endswith('.pdf'):
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""
    else:
        return ""

@app.route('/', methods=['GET', 'POST'])
def index():
    generated_resume = None
    error = None

    if request.method == 'POST':
        try:
            job_description = request.form.get('jobDescription', '')
            base_resume = ""

            # Handle resume upload
            if 'resume' in request.files:
                file = request.files['resume']
                if file.filename:
                    base_resume = read_resume_file(file)

            if job_description:
                # Initialize GPT4All with the mini model
                llm = ChatOpenAI(model="gpt-4o-mini")

                # Create prompt based on whether a resume was uploaded
                if base_resume:
                    prompt = f"""Create a professional resume in markdown format based on this job description and existing resume:

Job Description:
{job_description}

Existing Resume:
{base_resume}

Instructions:
- Use the existing resume as a base but modify it to better match the job requirements
- Format the output in proper markdown
- Include relevant sections (Summary, Skills, Experience, Education)
- Keep it to one page
- Make it professional and well-organized
- Maintain truthful information from the original resume
"""
                else:
                    prompt = f"""Create a professional resume in markdown format based on this job description:

Instructions:
- Format the resume in proper markdown
- Include relevant sections (Summary, Skills, Experience, Education)
- Create content that matches the job requirements
- Keep it to ONE PAGE ONLY
- Make it professional and well-organized
- Use realistic but fictional details
DO NOT INCLUDE ANYTHING ELSE THAN THE RESUME
Ensure that there is no other text saying markdown or anything else


Do it in this format:

# Sujar Henry  
571-356-7863 | sujarhenry@gmail.com | [linkedin.com/in/sujar-henry](https://linkedin.com/in/sujar-henry) | [github.com/Sujar-Henry](https://github.com/Sujar-Henry)

---

## Summary  
Detail-oriented and innovative Software Engineer Intern proficient in designing and developing AI and machine learning solutions. Adept at optimizing workflows and enhancing document processing efficiency. Strong foundation in computer science complemented by hands-on experience in deep learning, cloud technologies, and software development methodologies. Passionate about leveraging technology to automate processes and improve operational efficiency.

---

## Skills  

- **Languages:** Python, Java, C++, HTML/CSS, JavaScript  
- **Libraries/Frameworks:** PyTorch, TensorFlow, LangChain, NumPy, pandas  
- **Cloud Technologies:** AWS Lambda, AWS S3, AWS API Gateway, AWS SES, Amazon DynamoDB  
- **Developer Tools:** Git, IBM Cloud  
- **Methodologies:** Agile Development, Machine Learning, Deep Learning  

---

## Experience  

**IBM**  
*Software Engineer Intern, Co-Op*  
_August 2024 – Present_  
- Developed an AI agent for automated Statement of Work (SOW) generation, leveraging Llama3 and Granite models in LangChain, reducing manual drafting time by 80%.
- Enhanced the SOW agent by implementing a structured scope section, leading to a 30% increase in document approval efficiency.
- Engineered AI-driven agents for workflow automation, utilizing LangChain and LangGraph, resulting in a 50% boost in operational efficiency.

*Software Engineer Intern*  
_May 2024 – August 2024_  
- Developed an Agent Workflow for converting legal documentation into OpenFisca code, achieving a 70% reduction in manual coding time.
- Integrated customizable LLM options, enabling users to select from WatsonX and OpenAI models, resulting in a 50% increase in OpenFisca code creation speed.
- Improved system performance with a versatile Logic and Code Checker, reducing workflow waiting times by 40%.

**Resilience Inc.**  
*Software Engineer Intern*  
_May 2023 – August 2023_  
- Developed a deep learning model in TensorFlow for emotion recognition, achieving 85% accuracy on a test set of audio samples.
- Engineered a Multi-Layer Perceptron (MLP) with a hidden layer size of 300, achieving 92% accuracy and strong performance metrics (F1-score: 0.82, AUC: 0.89).

---

## Education  

**The Catholic University of America**, Washington, DC  
*Bachelor of Science in Computer Science, Minor in Mathematics*  
_Expected May 2026_  
Relevant Coursework: Data Structures, Web Design & Programming, Object-Oriented Programming with Java  

---

## Projects  

**Speech Translator**  
*Technologies: Python, PyTorch, LibreSpeech*  
- Implemented a speech-to-text ML model utilizing Wav2Vec2 and the LibriSpeech dataset, achieving high accuracy in emotion recognition.
- Leveraged LibreTranslate API for multilingual text-to-speech conversion.

**Serverless Portfolio**  
*Technologies: AWS (Lambda, S3), HTML, CSS*  
- Optimized S3 storage, reducing costs by 25% and ensuring rapid access speeds.
- Implemented a REST API to enhance email capacity, increasing sending speed from 60 to 840 emails per minute.

---

## Affiliations  
- National Society of Black Engineers  
- ColorStack  
- INROADS  

"""
                # Generate the tailored resume
                response = llm.invoke(prompt)
                if hasattr(response, 'content'):
                    response = response.content
                # Remove any "markdown" text that might appear in the response
                response = response.replace("```markdown", "").replace("```", "")
                
                generated_resume = str(response) if response else ""

                # Save as Markdown
                with open("resume.md", "w", encoding='utf-8') as f:
                    f.write(generated_resume)

                # Create Word document
                doc = Document()
                lines = generated_resume.split('\n')
                for line in lines:
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('- '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        doc.add_paragraph(line)

                doc.save("resume.docx")

                # Convert to PDF
                try:
                    convert("resume.docx", "resume.pdf")
                except Exception as e:
                    print(f"Error converting to PDF: {e}")

        except Exception as e:
            error = str(e)
            print(f"Error: {e}")

    return render_template('index.html', 
                         generated_resume=generated_resume, 
                         error=error)

@app.route('/download/<format>')
def download_file(format):
    if format in ['pdf', 'docx', 'md']:
        filename = f"resume.{format}"
        if os.path.exists(filename):
            return send_file(filename, as_attachment=True)
    return "File not found", 404

if __name__ == '__main__':
    app.run(debug=True, port=8000)